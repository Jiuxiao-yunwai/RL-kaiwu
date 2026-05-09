from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]


REPORTS = [
    (
        ROOT / "exp1-gorge-work" / "docs" / "实验1实验报告.md",
        ROOT / "exp1-gorge-work" / "docs" / "实验1实验报告_Word版.docx",
    ),
    (
        ROOT / "exp2-back-to-the-realm" / "docs" / "实验2实验报告.md",
        ROOT / "exp2-back-to-the-realm" / "docs" / "实验2实验报告_Word版.docx",
    ),
]


ACCENT = RGBColor(31, 78, 121)
LIGHT_FILL = "D9EAF7"
HEADER_FILL = "1F4E79"
CODE_FILL = "F3F6F8"
BORDER = "A6A6A6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_pct=5000):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_pct))
    tbl_w.set(qn("w:type"), "pct")


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in [
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11.5, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(28)
    r = sub.add_run("强化学习实验报告")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(89, 89, 89)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = line.add_run("—")
    lr.font.size = Pt(28)
    lr.font.color.rgb = ACCENT

    doc.add_page_break()


def split_table_row(line):
    raw = line.strip().strip("|")
    return [cell.strip().replace("\\_", "_") for cell in raw.split("|")]


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def add_markdown_text(paragraph, text, bold=False):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        run_bold = bold
        content = part
        if part.startswith("`") and part.endswith("`"):
            content = part[1:-1]
            run = paragraph.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(79, 79, 79)
            continue
        if part.startswith("**") and part.endswith("**"):
            content = part[2:-2]
            run_bold = True
        run = paragraph.add_run(content)
        run.bold = run_bold


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.15
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
                run = para.add_run(value)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 0:
                    set_cell_shading(cell, "F7FBFD")
                add_markdown_text(para, value)
            if len(row) <= 3 and c_idx == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif value.replace(".", "", 1).replace("-", "", 1).isdigit() or value in {"-", "0.1", "0.9", "0.99"}:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Compact obvious short columns.
    cols = len(rows[0])
    if cols >= 3:
        for row in table.rows:
            set_cell_width(row.cells[0], 1800)
    doc.add_paragraph()


def add_code_block(doc, code_lines):
    for line in code_lines:
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.right_indent = Cm(0.2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(64, 64, 64)
        p_pr = p._p.get_or_add_pPr()
        shd = p_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            p_pr.append(shd)
        shd.set(qn("w:fill"), CODE_FILL)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def build_docx(md_path, out_path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = lines[0].lstrip("# ").strip() if lines else md_path.stem

    doc = Document()
    configure_document(doc)
    add_cover(doc, title)

    i = 0
    in_code = False
    code_lines = []
    pending_table = []
    list_counter = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            pending_table = [split_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                pending_table.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, pending_table)
            pending_table = []
            continue

        if not stripped:
            list_counter = 0
            i += 1
            continue

        if stripped.startswith("# "):
            # Cover already contains the report title.
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        elif re.match(r"^\d+\.\s+", stripped):
            match = re.match(r"^(\d+)\.\s+(.*)", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.first_line_indent = Cm(0)
            add_markdown_text(p, f"{match.group(1)}、{match.group(2)}")
        elif stripped.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.first_line_indent = Cm(0)
            add_markdown_text(p, f"• {stripped[2:]}")
        else:
            p = doc.add_paragraph()
            add_markdown_text(p, stripped)
        i += 1

    doc.save(out_path)


def main():
    for md_path, out_path in REPORTS:
        build_docx(md_path, out_path)
        print(out_path)


if __name__ == "__main__":
    main()
